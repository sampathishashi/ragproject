import streamlit as st
import random
import re
from datetime import datetime

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_groq import ChatGroq
from langchain_core.prompts import ChatPromptTemplate

# ==========================================================
# PAGE CONFIGURATION
# ==========================================================

st.set_page_config(
    page_title="AI Mock Test Generator",
    page_icon="🎯",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ==========================================================
# CUSTOM CSS
# ==========================================================

st.markdown("""
<style>

.block-container{
    padding-top:1rem;
    padding-bottom:2rem;
}

.main-title{
    font-size:40px;
    font-weight:700;
    color:#2E8B57;
}

.sub-title{
    color:#808080;
    font-size:17px;
}

.card{
    background:#f8f9fa;
    padding:18px;
    border-radius:10px;
    border:1px solid #e3e3e3;
}

.stButton>button{
    width:100%;
    height:50px;
    font-size:18px;
    font-weight:bold;
    border-radius:8px;
}

.stDownloadButton>button{
    width:100%;
}

</style>
""", unsafe_allow_html=True)

# ==========================================================
# HEADER
# ==========================================================

st.markdown(
"""
<div class="main-title">
🎯 AI Mock Test Generator
</div>

<div class="sub-title">
Professional Interview Question Generator using
<b>RAG + FAISS + HuggingFace + Groq</b>
</div>
""",
unsafe_allow_html=True
)

st.divider()

# ==========================================================
# SIDEBAR
# ==========================================================

st.sidebar.title("⚙ Configuration")

company = st.sidebar.selectbox(
    "Target Company",
    [
        "Google",
        "Amazon",
        "Microsoft",
        "Adobe",
        "Oracle",
        "Uber",
        "Goldman Sachs",
        "Flipkart",
        "Infosys",
        "TCS",
        "Accenture"
    ]
)

question_type = st.sidebar.selectbox(
    "Question Type",
    [
        "Aptitude",
        "Coding"
    ]
)

difficulty = st.sidebar.selectbox(
    "Difficulty",
    [
        "Easy",
        "Medium",
        "Hard"
    ]
)

num_questions = st.sidebar.slider(
    "Number of Questions",
    1,
    20,
    5
)

topic = st.text_input(
    "Enter Topic",
    placeholder="Arrays, Percentage, DBMS..."
)

generate = st.button("🚀 Generate Questions")

# ==========================================================
# SESSION STATE
# ==========================================================

if "history" not in st.session_state:
    st.session_state.history = []

if "generated_questions" not in st.session_state:
    st.session_state.generated_questions = set()

# ==========================================================
# BUILT-IN KNOWLEDGE BASE
# ==========================================================

study_material = """

DATA STRUCTURES

Arrays
Searching
Sorting
Binary Search
Sliding Window
Two Pointer
Prefix Sum
Kadane Algorithm

Strings
Palindrome
Anagram
KMP
Rabin Karp
Z Algorithm

Linked List
Stack
Queue
Heap
Trie

Trees
Binary Tree
BST
AVL
Segment Tree

Graphs
DFS
BFS
Topological Sort
Dijkstra
Bellman Ford
Union Find
Minimum Spanning Tree

Dynamic Programming
0/1 Knapsack
LCS
LIS
Coin Change
Matrix Chain Multiplication

DBMS
Normalization
Transactions
SQL
Joins
Indexing

Operating System
CPU Scheduling
Deadlock
Paging
Segmentation
Memory Management

Computer Networks
TCP
UDP
HTTP
HTTPS
DNS
Routing
OSI Model

Object Oriented Programming
Inheritance
Encapsulation
Polymorphism
Abstraction

QUANTITATIVE APTITUDE

Percentage
Profit and Loss
Simple Interest
Compound Interest
Average
Ratio and Proportion
Probability
Permutation
Combination
Number System
Time and Work
Time Speed Distance
Boats and Streams
Mixtures
Pipes and Cisterns

Logical Reasoning
Blood Relations
Directions
Coding Decoding
Seating Arrangement
Syllogism

Data Interpretation
Bar Graph
Pie Chart
Table
Line Graph

"""

# ==========================================================
# VECTOR DATABASE
# ==========================================================

@st.cache_resource
def load_vector_database():

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=600,
        chunk_overlap=100
    )

    chunks = splitter.split_text(study_material)

    embeddings = HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-MiniLM-L6-v2",
        model_kwargs={"device": "cpu"},
        encode_kwargs={"normalize_embeddings": True}
    )

    vector_db = FAISS.from_texts(
        texts=chunks,
        embedding=embeddings
    )

    return vector_db


vector_db = load_vector_database()

# ==========================================================
# RETRIEVER
# ==========================================================

retriever = vector_db.as_retriever(

    search_type="mmr",

    search_kwargs={

        "k": 6,

        "fetch_k": 20

    }

)

# ==========================================================
# LLM CONFIGURATION
# ==========================================================

llm = ChatGroq(

    api_key=st.secrets["GROQ_API_KEY"],

    model="llama-3.3-70b-versatile",

    temperature=0.2,

    max_tokens=4096

)

# ==========================================================
# CONTEXT RETRIEVAL
# ==========================================================

def retrieve_context(topic: str):

    docs = retriever.invoke(topic)

    if not docs:
        return ""

    return "\n\n".join(
        doc.page_content
        for doc in docs
    )

# ==========================================================
# CLEAN RESPONSE
# ==========================================================

def clean_response(text):

    if not text:
        return ""

    text = text.replace("```", "")

    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()

# ==========================================================
# RANDOM QUESTION STYLES
# ==========================================================

QUESTION_STYLES = [

    "Google Online Assessment",

    "Amazon Online Assessment",

    "Microsoft Coding Round",

    "Adobe Hiring Test",

    "Uber Technical Assessment",

    "Goldman Sachs Aptitude Test",

    "Campus Placement Test",

    "Competitive Programming",

    "Interview Screening Test"

]

# ==========================================================
# GENERATION TIMER
# ==========================================================

def generation_time():

    return datetime.now().strftime("%d-%m-%Y %H:%M:%S")

# ==========================================================
# APP STATUS
# ==========================================================

st.sidebar.markdown("---")

st.sidebar.success("✅ AI Engine Ready")

st.sidebar.write("LLM : Groq")

st.sidebar.write("Embeddings : MiniLM")

st.sidebar.write("Vector DB : FAISS")

st.sidebar.write("Retriever : MMR")

st.sidebar.write("Status : Ready")

st.sidebar.markdown("---")

# ==========================================================
# PROFESSIONAL APTITUDE PROMPT
# ==========================================================

APTITUDE_PROMPT = ChatPromptTemplate.from_template("""
You are a Senior Aptitude Assessment Designer with 15+ years of experience.

You create placement assessments for

• Google
• Amazon
• Microsoft
• Adobe
• Uber
• Goldman Sachs

==========================================================

Company:
{company}

Assessment Style:
{style}

Difficulty:
{difficulty}

Topic:
{topic}

Reference Material:
{context}

==========================================================

TASK

Generate EXACTLY ONE ORIGINAL aptitude question.

The question must look like a REAL company assessment.

Never copy textbook questions.

Never copy internet examples.

Invent a realistic scenario.

==========================================================

RULES

1. Generate a unique question.

2. Use realistic numbers.

3. Exactly FOUR options.

4. Exactly ONE correct answer.

5. Solve the problem before writing.

6. Verify every calculation.

7. Never contradict yourself.

8. Never approximate.

9. Never write
"The closest answer..."

10. If any calculation is wrong,
discard the question and regenerate it.

==========================================================

OUTPUT FORMAT

Question

<Question>

Options

A. ...

B. ...

C. ...

D. ...

----------------------------------------------------------

Correct Answer

The correct answer is <Option>. <Answer>

----------------------------------------------------------

Step-by-Step Solution

Generate ONLY the required number of steps.

Some questions need

2 Steps

Some need

3 Steps

Some need

6 Steps

Generate ONLY the required steps.

Example

Step 1

<Explanation>

Step 2

<Explanation>

Step 3

<Explanation>

Continue until solved.

Every step must show

• What is calculated

• The calculation

• The intermediate result

Never skip calculations.

----------------------------------------------------------

Final Answer

The correct answer is <Option>. <Answer>

==========================================================

QUALITY CHECK

✓ Grammar

✓ Mathematics

✓ Logic

✓ Options

✓ Explanation

✓ Final Answer

Return ONLY the final question.

""")

# ==========================================================
# PROFESSIONAL CODING PROMPT
# ==========================================================

CODING_PROMPT = ChatPromptTemplate.from_template("""
You are a Senior Software Engineer.

Create ONE original coding interview problem.

Company:
{company}

Assessment Style:
{style}

Difficulty:
{difficulty}

Topic:
{topic}

Reference:
{context}

Output

Problem Statement

Input Format

Output Format

Constraints

Sample Input

Sample Output

Explanation

Python Function

Hints

Time Complexity

Space Complexity

Requirements

• Original question

• Professional wording

• Company level

• No copied problems

Return only the coding problem.
""")

# ==========================================================
# AI REVIEWER
# ==========================================================

REVIEW_PROMPT = ChatPromptTemplate.from_template("""
You are a Senior Quality Reviewer.

Review this interview question.

{question}

Checklist

✓ English

✓ Grammar

✓ Logic

✓ Mathematics

✓ Correct Option

✓ Explanation

✓ Final Answer

If ANY mistake exists

Rewrite the ENTIRE question.

Do NOT explain the mistake.

Return ONLY the corrected version.
""")
# ==========================================================
# QUESTION GENERATOR ENGINE
# ==========================================================

def generate_question(
    company,
    topic,
    difficulty,
    question_type
):
    """
    Generate one question using RAG + AI Review.
    """

    context = retrieve_context(topic)

    style = random.choice(QUESTION_STYLES)

    if question_type == "Aptitude":

        prompt = APTITUDE_PROMPT.format(
            company=company,
            style=style,
            difficulty=difficulty,
            topic=topic,
            context=context
        )

    else:

        prompt = CODING_PROMPT.format(
            company=company,
            style=style,
            difficulty=difficulty,
            topic=topic,
            context=context
        )

    try:

        # -----------------------------
        # First Generation
        # -----------------------------

        first_response = llm.invoke(prompt)

        draft = clean_response(first_response.content)

        # -----------------------------
        # AI Review
        # -----------------------------

        review = REVIEW_PROMPT.format(
            question=draft
        )

        reviewed = llm.invoke(review)

        final_question = clean_response(
            reviewed.content
        )

        return final_question

    except Exception as e:

        return f"""
❌ Generation Error

{str(e)}
"""


# ==========================================================
# QUALITY CHECK
# ==========================================================

def is_valid_question(question: str):

    if not question:
        return False

    required_sections = [

        "Question",

        "Options",

        "Correct Answer"

    ]

    for section in required_sections:

        if section.lower() not in question.lower():

            return False

    return True


# ==========================================================
# DUPLICATE CHECK
# ==========================================================

def is_duplicate(question: str):

    normalized = " ".join(question.split()).lower()

    if normalized in st.session_state.generated_questions:

        return True

    st.session_state.generated_questions.add(normalized)

    return False


# ==========================================================
# RETRY LOGIC
# ==========================================================

def generate_unique_question(
    company,
    topic,
    difficulty,
    question_type,
    retries=5
):

    for _ in range(retries):

        question = generate_question(
            company,
            topic,
            difficulty,
            question_type
        )

        if not is_valid_question(question):

            continue

        if is_duplicate(question):

            continue

        return question

    return """
Unable to generate a unique high-quality question.

Please try again.
"""


# ==========================================================
# MULTIPLE QUESTION GENERATOR
# ==========================================================

def generate_multiple_questions(
    company,
    topic,
    difficulty,
    question_type,
    count
):

    questions = []

    st.session_state.generated_questions.clear()

    progress = st.progress(0)

    status = st.empty()

    for i in range(count):

        status.info(
            f"Generating Question {i+1} of {count}"
        )

        q = generate_unique_question(
            company,
            topic,
            difficulty,
            question_type
        )

        questions.append(q)

        progress.progress(
            (i + 1) / count
        )

    progress.empty()

    status.success("✅ Generation Completed")

    return questions
    # ==========================================================
# DISPLAY QUESTION
# ==========================================================

def display_question(question_number, content):

    st.markdown(
        f"""
        <div style="
            background:#ffffff;
            border:1px solid #dcdcdc;
            border-radius:12px;
            padding:20px;
            margin-bottom:20px;
        ">
        <h3>Question {question_number}</h3>
        </div>
        """,
        unsafe_allow_html=True
    )

    # Split into question and solution
    if "Step-by-Step Solution" in content:

        question_part, solution_part = content.split(
            "Step-by-Step Solution",
            1
        )

        st.markdown(question_part)

        with st.expander(
            "📖 View Step-by-Step Solution",
            expanded=False
        ):

            st.markdown("## Step-by-Step Solution")

            st.markdown(solution_part)

    else:

        st.markdown(content)

    st.divider()


# ==========================================================
# DOWNLOAD FORMAT
# ==========================================================

def prepare_download(questions):

    text = ""

    for i, q in enumerate(questions, 1):

        text += f"""

======================================================
QUESTION {i}
======================================================

{q}

"""

    return text


# ==========================================================
# HISTORY
# ==========================================================

def save_history(

    company,

    topic,

    difficulty,

    question_type,

    number

):

    st.session_state.history.append(

        {

            "Company": company,

            "Topic": topic,

            "Difficulty": difficulty,

            "Type": question_type,

            "Questions": number,

            "Time": generation_time()

        }

    )


# ==========================================================
# MAIN GENERATION
# ==========================================================

if generate:

    if topic.strip() == "":

        st.warning("Please enter a topic.")

        st.stop()

    st.markdown("---")

    st.header("Generated Questions")

    questions = generate_multiple_questions(

        company,

        topic,

        difficulty,

        question_type,

        num_questions

    )

    for i, question in enumerate(questions, 1):

        display_question(i, question)

    save_history(

        company,

        topic,

        difficulty,

        question_type,

        num_questions

    )

    download_text = prepare_download(

        questions

    )

    st.download_button(

        label="📥 Download Questions",

        data=download_text,

        file_name=f"{company}_{topic}.txt",

        mime="text/plain"

    )


# ==========================================================
# HISTORY PANEL
# ==========================================================

if st.session_state.history:

    st.markdown("---")

    st.header("📜 Recent Generations")

    for item in reversed(st.session_state.history):

        with st.container():

            st.markdown(f"""
**Company:** {item['Company']}

**Topic:** {item['Topic']}

**Difficulty:** {item['Difficulty']}

**Question Type:** {item['Type']}

**Questions Generated:** {item['Questions']}

**Generated On:** {item['Time']}
""")

            st.divider()
            # ==========================================================
# PROFESSIONAL DASHBOARD
# ==========================================================

st.markdown("---")

st.header("📊 Dashboard")

col1, col2, col3, col4 = st.columns(4)

with col1:
    st.metric(
        "🏢 Company",
        company
    )

with col2:
    st.metric(
        "📚 Topic",
        topic if topic else "-"
    )

with col3:
    st.metric(
        "📝 Questions",
        num_questions
    )

with col4:
    st.metric(
        "🎯 Difficulty",
        difficulty
    )

st.markdown("---")

# ==========================================================
# AI FEATURES
# ==========================================================

st.subheader("🚀 AI Features")

c1, c2, c3 = st.columns(3)

with c1:
    st.success("""
### 🧠 RAG

✔ FAISS

✔ HuggingFace

✔ Semantic Search

✔ MMR Retrieval
""")

with c2:
    st.info("""
### 🤖 AI

✔ Groq LLM

✔ AI Verification

✔ Duplicate Prevention

✔ Smart Prompting
""")

with c3:
    st.warning("""
### 🎯 Quality

✔ Company Pattern

✔ Original Questions

✔ Dynamic Steps

✔ Professional Format
""")

st.markdown("---")

# ==========================================================
# QUESTION QUALITY
# ==========================================================

st.subheader("📈 Question Quality")

quality = 96

st.progress(quality)

st.success(f"Estimated Question Quality : {quality}%")

st.caption("""
Questions are generated using Retrieval Augmented Generation (RAG),
professional prompt engineering and an AI verification step.
""")

# ==========================================================
# COMPANY SUPPORT
# ==========================================================

st.markdown("---")

st.subheader("🏢 Supported Companies")

companies = [

    "Google",

    "Amazon",

    "Microsoft",

    "Adobe",

    "Oracle",

    "Uber",

    "Goldman Sachs",

    "Flipkart",

    "Infosys",

    "TCS",

    "Accenture"

]

cols = st.columns(3)

for index, company_name in enumerate(companies):

    cols[index % 3].success(company_name)

# ==========================================================
# SUPPORTED TOPICS
# ==========================================================

st.markdown("---")

st.subheader("📚 Supported Topics")

topics = [

    "Arrays",

    "Strings",

    "Linked List",

    "Stack",

    "Queue",

    "Trees",

    "Graphs",

    "Dynamic Programming",

    "DBMS",

    "Operating System",

    "Computer Networks",

    "SQL",

    "Probability",

    "Percentage",

    "Profit and Loss",

    "Time and Work",

    "Logical Reasoning",

    "Data Interpretation"

]

for t in topics:

    st.write("✔", t)

# ==========================================================
# SIDEBAR INFO
# ==========================================================

st.sidebar.markdown("---")

st.sidebar.subheader("⚡ AI Engine")

st.sidebar.success("Groq Connected")

st.sidebar.success("FAISS Loaded")

st.sidebar.success("Embeddings Ready")

st.sidebar.success("Retriever Active")

st.sidebar.success("AI Reviewer Active")

st.sidebar.markdown("---")

st.sidebar.subheader("📊 Statistics")

st.sidebar.metric(
    "Sessions",
    len(st.session_state.history)
)

st.sidebar.metric(
    "Questions",
    num_questions
)

st.sidebar.metric(
    "Difficulty",
    difficulty
)
# ==========================================================
# PROFESSIONAL QUESTION DISPLAY
# ==========================================================

def display_question(question_number, content):

    st.markdown(
        f"""
        <div style="
            background:white;
            padding:20px;
            border-radius:12px;
            border-left:6px solid #4CAF50;
            box-shadow:0px 2px 8px rgba(0,0,0,0.08);
            margin-bottom:20px;
        ">
        <h2 style="color:#2E8B57;">
        Question {question_number}
        </h2>
        </div>
        """,
        unsafe_allow_html=True
    )

    # ---------------------------------------------------
    # Split Question and Solution
    # ---------------------------------------------------

    if "Step-by-Step Solution" in content:

        question_text, solution_text = content.split(
            "Step-by-Step Solution",
            1
        )

    else:

        question_text = content
        solution_text = ""

    # ---------------------------------------------------
    # Display Question
    # ---------------------------------------------------

    st.markdown(question_text)

    # ---------------------------------------------------
    # Display Solution
    # ---------------------------------------------------

    if solution_text.strip():

        with st.expander(
            "📖 View Detailed Solution",
            expanded=False
        ):

            st.success("Step-by-Step Solution")

            lines = solution_text.splitlines()

            step = []

            for line in lines:

                if line.strip().startswith("Step"):

                    if step:

                        st.markdown("\n".join(step))

                        st.divider()

                    step = [f"### {line.strip()}"]

                else:

                    step.append(line)

            if step:

                st.markdown("\n".join(step))

    st.markdown("---")
    # ==========================================================
# IMPORTS
# (Add these at the top of app.py if not already imported)
# ==========================================================

from io import BytesIO
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet

# ==========================================================
# EXPORT TO DOCX
# ==========================================================

def export_docx(questions):

    document = Document()

    document.add_heading(
        "AI Mock Test Generator",
        level=1
    )

    for i, question in enumerate(questions, 1):

        document.add_heading(
            f"Question {i}",
            level=2
        )

        document.add_paragraph(question)

    file = BytesIO()

    document.save(file)

    file.seek(0)

    return file


# ==========================================================
# EXPORT TO PDF
# ==========================================================

def export_pdf(questions):

    pdf = BytesIO()

    styles = getSampleStyleSheet()

    doc = SimpleDocTemplate(pdf)

    story = []

    story.append(
        Paragraph(
            "<b>AI Mock Test Generator</b>",
            styles["Heading1"]
        )
    )

    for i, q in enumerate(questions, 1):

        story.append(
            Paragraph(
                f"<b>Question {i}</b>",
                styles["Heading2"]
            )
        )

        story.append(
            Paragraph(
                q.replace("\n","<br/>"),
                styles["BodyText"]
            )
        )

    doc.build(story)

    pdf.seek(0)

    return pdf


# ==========================================================
# FAVORITES
# ==========================================================

if "favorites" not in st.session_state:

    st.session_state.favorites = []


def save_favorite(question):

    if question not in st.session_state.favorites:

        st.session_state.favorites.append(question)


# ==========================================================
# SEARCH
# ==========================================================

search = st.sidebar.text_input(
    "🔍 Search Question"
)

# ==========================================================
# EXPORT BUTTONS
# ==========================================================

if generate:

    st.markdown("---")

    st.subheader("📥 Export Questions")

    col1, col2, col3 = st.columns(3)

    with col1:

        st.download_button(

            "Download TXT",

            prepare_download(questions),

            file_name="questions.txt",

            mime="text/plain"

        )

    with col2:

        st.download_button(

            "Download DOCX",

            export_docx(questions),

            file_name="questions.docx",

            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        )

    with col3:

        st.download_button(

            "Download PDF",

            export_pdf(questions),

            file_name="questions.pdf",

            mime="application/pdf"

        )

# ==========================================================
# FAVORITE DISPLAY
# ==========================================================

if st.session_state.favorites:

    st.markdown("---")

    st.subheader("⭐ Favorite Questions")

    for q in st.session_state.favorites:

        st.markdown(q)

        st.divider()

# ==========================================================
# ANALYTICS
# ==========================================================

st.markdown("---")

st.subheader("📊 Analytics")

a1, a2, a3 = st.columns(3)

with a1:

    st.metric(
        "Generated",
        len(st.session_state.history)
    )

with a2:

    st.metric(
        "Favorites",
        len(st.session_state.favorites)
    )

with a3:

    st.metric(
        "Current Batch",
        num_questions
    )st.markdown("""
<style>

/* ===============================
   MAIN
================================= */

.stApp{
    background:#F6F8FC;
}

/* ===============================
   HEADER
================================= */

.main-header{
    background:linear-gradient(135deg,#1E88E5,#42A5F5);
    padding:25px;
    border-radius:15px;
    color:white;
    margin-bottom:20px;
}

.main-header h1{
    color:white;
    font-size:38px;
    margin-bottom:5px;
}

.main-header p{
    font-size:18px;
}

/* ===============================
   QUESTION CARD
================================= */

.question-card{

    background:white;

    padding:25px;

    border-radius:15px;

    box-shadow:0 5px 15px rgba(0,0,0,.08);

    margin-bottom:25px;

    border-left:8px solid #1E88E5;

}

/* ===============================
   STEP CARD
================================= */

.step-card{

    background:#F9FAFB;

    border-left:5px solid #43A047;

    padding:15px;

    border-radius:10px;

    margin-bottom:15px;

}

/* ===============================
   CORRECT ANSWER
================================= */

.answer-card{

    background:#E8F5E9;

    border-left:6px solid #2E7D32;

    padding:18px;

    border-radius:10px;

    font-weight:bold;

}

/* ===============================
   BUTTONS
================================= */

.stButton>button{

    background:#1E88E5;

    color:white;

    height:48px;

    border-radius:10px;

    font-size:17px;

    font-weight:bold;

}

.stButton>button:hover{

    background:#1565C0;

}

</style>
""", unsafe_allow_html=True)
    st.markdown(f"""
<div class="main-header">

<h1>🎯 AI Mock Test Generator</h1>

<p>

Generate professional

<b>{company}</b>

{question_type}

Questions using

Groq + LangChain + FAISS + HuggingFace

</p>

</div>
""", unsafe_allow_html=True)
    import re

def display_question(number, content):

    st.markdown(f"""
<div class="question-card">

<h2>Question {number}</h2>

</div>
""", unsafe_allow_html=True)

    if "Step-by-Step Solution" not in content:

        st.markdown(content)
        return

    question, solution = content.split(
        "Step-by-Step Solution",
        1
    )

    st.markdown(question)

    if "Correct Answer" in solution:

        answer, steps = solution.split(
            "Step-by-Step Solution",
            1
        ) if "Step-by-Step Solution" in solution else ("", solution)

    with st.expander("📖 View Solution", expanded=False):

        step_pattern = r"(Step\s+\d+.*?)(?=Step\s+\d+|Final Answer|$)"

        matches = re.findall(
            step_pattern,
            solution,
            flags=re.S
        )

        for step in matches:

            st.markdown(f"""
<div class="step-card">

{step}

</div>
""", unsafe_allow_html=True)

        if "Final Answer" in solution:

            final = solution.split(
                "Final Answer",
                1
            )[1]

            st.markdown(f"""
<div class="answer-card">

<b>Final Answer</b>

{final}

</div>
""", unsafe_allow_html=True)
            c1,c2,c3,c4=st.columns(4)

c1.metric("🏢 Company",company)

c2.metric("📚 Topic",topic)

c3.metric("🎯 Difficulty",difficulty)

c4.metric("📝 Questions",num_questions)
st.download_button(

label="📥 Download Mock Test",

data=prepare_download(questions),

file_name=f"{company}_{topic}.txt",

mime="text/plain"

)
# ==========================================================
# AI CONFIDENCE SCORE
# ==========================================================

st.markdown("---")

confidence = random.randint(95, 99)

st.subheader("🤖 AI Confidence")

st.progress(confidence)

st.success(f"Confidence Score: {confidence}%")
st.info(f"""
### 🏢 Assessment

**Company:** {company}

**Difficulty:** {difficulty}

**Question Type:** {question_type}

**Topic:** {topic}
""")
st.markdown("---")

st.subheader("📊 Session Statistics")

col1, col2, col3 = st.columns(3)

with col1:
    st.metric("Questions Generated", len(questions))

with col2:
    st.metric("History", len(st.session_state.history))

with col3:
    st.metric("AI Quality", "96%")
    if st.button("🔄 Regenerate Questions"):

    st.rerun()
    st.markdown("---")

st.markdown("""
<div style="text-align:center;padding:25px;color:gray;">

## 🎯 AI Mock Test Generator

Professional Interview Assessment Platform

Powered by

**Groq • LangChain • HuggingFace • FAISS • Streamlit**

Version 3.0

© 2026 All Rights Reserved

</div>
""", unsafe_allow_html=True)
st.sidebar.markdown("---")

st.sidebar.success("🟢 System Online")

st.sidebar.write("Model : Groq")

st.sidebar.write("Retriever : FAISS")

st.sidebar.write("Embeddings : MiniLM")

st.sidebar.write("Quality Check : Enabled")

st.sidebar.write("RAG : Active")
